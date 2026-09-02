"""Resume file -> plain text. PDF via PyMuPDF, DOCX via python-docx."""

from __future__ import annotations

import io
import logging
import re
from pathlib import Path

log = logging.getLogger("proofscreen.ingest")

SUPPORTED = {".pdf", ".docx", ".txt", ".md"}
MAX_BYTES = 10 * 1024 * 1024


class UnsupportedResume(ValueError):
    """Wrong file type, empty file, or unreadable document."""


_WS = re.compile(r"[ \t ]+")
_BLANKS = re.compile(r"\n{3,}")


def normalise(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = _WS.sub(" ", text)
    text = "\n".join(line.strip() for line in text.split("\n"))
    return _BLANKS.sub("\n\n", text).strip()


def _from_pdf(data: bytes) -> str:
    try:
        import pymupdf as fitz          # pymupdf >= 1.24.3
    except ImportError:
        try:
            import fitz                 # older releases only expose `fitz`
        except ImportError as exc:  # pragma: no cover
            raise UnsupportedResume("PDF support needs pymupdf installed") from exc
    try:
        with fitz.open(stream=data, filetype="pdf") as doc:
            return "\n".join(page.get_text("text") for page in doc)
    except Exception as exc:  # noqa: BLE001
        raise UnsupportedResume(f"could not read PDF: {exc}") from exc


def _from_docx(data: bytes) -> str:
    try:
        import docx  # python-docx
    except ImportError as exc:  # pragma: no cover
        raise UnsupportedResume("DOCX support needs python-docx installed") from exc
    try:
        document = docx.Document(io.BytesIO(data))
    except Exception as exc:  # noqa: BLE001
        raise UnsupportedResume(f"could not read DOCX: {exc}") from exc

    parts = [p.text for p in document.paragraphs]
    for table in document.tables:            # resumes love tables
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def extract_text(filename: str, data: bytes) -> str:
    """Main entry point. Raises UnsupportedResume with a message fit for a 400."""
    if not data:
        raise UnsupportedResume("file is empty")
    if len(data) > MAX_BYTES:
        raise UnsupportedResume(f"file is larger than {MAX_BYTES // (1024 * 1024)}MB")

    suffix = Path(filename or "").suffix.lower()
    if suffix not in SUPPORTED:
        raise UnsupportedResume(
            f"unsupported file type '{suffix or filename}'. "
            f"Accepted: {', '.join(sorted(SUPPORTED))}"
        )

    if suffix == ".pdf":
        text = _from_pdf(data)
    elif suffix == ".docx":
        text = _from_docx(data)
    else:
        text = data.decode("utf-8", errors="replace")

    text = normalise(text)
    if len(text) < 80:
        raise UnsupportedResume(
            "extracted almost no text — the file may be a scan. "
            "Paste the resume text into POST /api/candidates/text instead."
        )
    log.info("parsed %s -> %d chars", filename, len(text))
    return text
